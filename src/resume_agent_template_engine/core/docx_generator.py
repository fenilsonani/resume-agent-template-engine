from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE # For defining styles if needed
from typing import Dict, Any, List

def _add_optional_section(document: Document, title: str, content_list: List[Dict[str, Any]], item_formatter_func):
    """Helper to add a section if content_list is not empty."""
    if content_list:
        document.add_heading(title, level=1)
        for item_data in content_list:
            item_formatter_func(document, item_data)
        document.add_paragraph() # Add space after the section

def _format_experience_item(document: Document, exp: Dict[str, Any]):
    """Formats and adds a single experience item to the document."""
    document.add_heading(exp.get("company", "N/A Company"), level=2)
    
    title_dates = f"{exp.get('title', 'N/A Title')}"
    dates = []
    if exp.get("startDate"):
        dates.append(exp.get("startDate"))
    if exp.get("endDate"):
        dates.append(exp.get("endDate"))
    if dates:
        title_dates += f" ({' – '.join(dates)})"
    
    p_title = document.add_paragraph(title_dates)
    if exp.get("location"):
        p_title.add_run(f" | {exp.get('location')}").italic = True
    
    details = exp.get("details", [])
    if isinstance(details, list):
        for detail in details:
            document.add_paragraph(str(detail), style='ListBullet')
    elif isinstance(details, str): # Handle if details is a single string
         document.add_paragraph(details, style='ListBullet')
    document.add_paragraph()

def _format_education_item(document: Document, edu: Dict[str, Any]):
    """Formats and adds a single education item to the document."""
    document.add_heading(edu.get("institution", "N/A Institution"), level=2)
    
    degree_field = edu.get("degree", "N/A Degree")
    if edu.get("field"):
        degree_field += f" in {edu.get('field')}"
    
    p_degree = document.add_paragraph()
    p_degree.add_run(degree_field).bold = True
    
    if edu.get("date") or edu.get("graduationYear"):
        date_str = edu.get("date", edu.get("graduationYear", ""))
        p_degree.add_run(f" ({date_str})")

    details = edu.get("details", [])
    if isinstance(details, list):
        for detail in details:
            # Could be bullet points or a single paragraph
            document.add_paragraph(str(detail), style='ListBullet') 
    elif isinstance(details, str):
        document.add_paragraph(details, style='ListBullet')
    document.add_paragraph()


def _format_project_item(document: Document, proj: Dict[str, Any]):
    """Formats and adds a single project item to the document."""
    document.add_heading(proj.get("name", "N/A Project Name"), level=2)
    
    if proj.get("technologies") and isinstance(proj.get("technologies"), list):
        tech_paragraph = document.add_paragraph()
        tech_paragraph.add_run("Technologies: ").italic = True
        tech_paragraph.add_run(", ".join(proj.get("technologies")))
        
    if proj.get("description"):
        document.add_paragraph(str(proj.get("description")))
    
    # If projects have 'details' or 'achievements' similar to experience:
    details = proj.get("details", proj.get("achievements", [])) # Check for both
    if isinstance(details, list):
        for detail in details:
            document.add_paragraph(str(detail), style='ListBullet')
    elif isinstance(details, str):
         document.add_paragraph(details, style='ListBullet')
    document.add_paragraph()


def generate_resume_docx(resume_data_dict: Dict[str, Any], output_path: str) -> str:
    """
    Generates a DOCX resume from a dictionary representing ResumeData.

    Args:
        resume_data_dict (Dict[str, Any]): A dictionary containing resume information.
        output_path (str): The path where the generated DOCX should be saved.

    Returns:
        str: The path to the generated DOCX file.
    """
    document = Document()
    document.styles['Normal'].font.name = 'Calibri'
    document.styles['Normal'].font.size = Pt(11)

    # --- Personal Info ---
    pi = resume_data_dict.get("personalInfo", {})
    if pi.get("name"):
        name_para = document.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(pi.get("name"))
        name_run.font.size = Pt(24)
        name_run.font.bold = True

    contact_lines = []
    if pi.get("email"):
        contact_lines.append(pi.get("email"))
    if pi.get("phone"):
        contact_lines.append(pi.get("phone"))
    if pi.get("location"):
        contact_lines.append(pi.get("location"))
    if pi.get("website"): # Assuming 'website' is the URL
        contact_lines.append(pi.get("website"))
    if pi.get("linkedin"): # Assuming 'linkedin' is the URL
        contact_lines.append(pi.get("linkedin"))
    
    if contact_lines:
        contact_para = document.add_paragraph(" | ".join(contact_lines))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(18)

    # --- Professional Summary ---
    summary = resume_data_dict.get("professional_summary")
    if summary:
        document.add_heading("Professional Summary", level=1)
        document.add_paragraph(str(summary))
        document.add_paragraph() # Add space

    # --- Experience ---
    _add_optional_section(document, "Experience", resume_data_dict.get("experience", []), _format_experience_item)

    # --- Education ---
    _add_optional_section(document, "Education", resume_data_dict.get("education", []), _format_education_item)
    
    # --- Projects ---
    _add_optional_section(document, "Projects", resume_data_dict.get("projects", []), _format_project_item)

    # --- Skills ---
    skills_data = resume_data_dict.get("technologies_and_skills")
    if skills_data:
        document.add_heading("Skills", level=1)
        if isinstance(skills_data, list):
            # Simple list of skills
            for skill_item in skills_data:
                 document.add_paragraph(str(skill_item), style='ListBullet')
        elif isinstance(skills_data, dict):
            # Categorized skills
            for category, skills_list in skills_data.items():
                if isinstance(skills_list, list):
                    document.add_paragraph(f"{category}: {', '.join(map(str, skills_list))}")
                else: # If skills_list is a single string under category
                    document.add_paragraph(f"{category}: {str(skills_list)}")
        document.add_paragraph()

    # --- Articles & Publications (Optional) ---
    pubs = resume_data_dict.get("articles_and_publications", [])
    if pubs:
        document.add_heading("Articles & Publications", level=1)
        for pub in pubs:
            pub_title = pub.get("title", "N/A Title")
            pub_details = []
            if pub.get("publisher"): pub_details.append(pub.get("publisher"))
            if pub.get("date"): pub_details.append(pub.get("date"))
            
            p = document.add_paragraph()
            p.add_run(pub_title).bold = True
            if pub_details:
                p.add_run(f" ({', '.join(pub_details)})")
            if pub.get("url"):
                 # Basic URL display, python-docx doesn't have simple hyperlink add to run.
                 # For actual hyperlinks, more complex XML manipulation is needed or using a template.
                p.add_run(f" - Available at: {pub.get('url')}").italic = True
        document.add_paragraph()

    # --- Certifications (Optional) ---
    certs = resume_data_dict.get("certifications", [])
    if certs:
        document.add_heading("Certifications", level=1)
        for cert in certs:
            cert_name = cert.get("name", "N/A Certification")
            cert_details = []
            if cert.get("issuer"): cert_details.append(cert.get("issuer"))
            if cert.get("date"): cert_details.append(cert.get("date"))

            p = document.add_paragraph()
            p.add_run(cert_name).bold = True
            if cert_details:
                p.add_run(f" ({', '.join(cert_details)})")
            if cert.get("url"):
                p.add_run(f" - Verify at: {cert.get('url')}").italic = True
        document.add_paragraph()

    try:
        document.save(output_path)
    except Exception as e:
        # Basic error handling, could be more specific if certain python-docx errors are common
        raise IOError(f"Error saving DOCX document to {output_path}: {e}") from e
    
    return output_path
```
