import pytest
import os
from docx import Document as DocxDocument # Alias to avoid confusion with other Document classes
from resume_agent_template_engine.core.docx_generator import generate_resume_docx
from typing import Dict, Any

# Minimal valid resume data for DOCX generation tests
MINIMAL_RESUME_DATA_DICT: Dict[str, Any] = {
    "personalInfo": {
        "name": "John Doe DOCX",
        "email": "john.docx@example.com",
        "phone": "123-456-7890",
        "location": "Testville, USA",
        "website": "johndocx.dev",
        "linkedin": "linkedin.com/in/johndocx"
    },
    "professional_summary": "A professional skilled in DOCX generation.",
    "experience": [
        {
            "title": "DOCX Specialist",
            "company": "Word Corp",
            "location": "Redmond, WA",
            "startDate": "2022-01",
            "endDate": "Present",
            "details": [
                "Created many DOCX files.",
                "Ensured DOCX quality and standards."
            ]
        }
    ],
    "education": [
        {
            "degree": "Master of Documents",
            "institution": "University of Files",
            "location": "Info City",
            "date": "2021",
            "field": "Digital Formatting"
        }
    ],
    "projects": [
        {
            "name": "DOCX Resume Generator Test",
            "technologies": ["python-docx", "pytest"],
            "description": "A project to test DOCX resume generation."
        }
    ],
    "technologies_and_skills": ["Microsoft Word", "python-docx", "Office Open XML"],
    "articles_and_publications": [ # Test optional section
        {
            "title": "The Art of DOCX",
            "publisher": "Word Monthly",
            "date": "2023",
            "url": "example.com/artofdocx"
        }
    ],
    "certifications": [ # Test optional section
        {
            "name": "Certified DOCX Professional",
            "issuer": "OpenXML Institute",
            "date": "2023",
            "url": "example.com/certprof"
        }
    ]
}

class TestDocxGenerator:

    def test_successful_docx_creation(self, tmp_path):
        """Test that generate_resume_docx runs without error and creates a file."""
        output_docx_file = tmp_path / "test_resume.docx"
        
        try:
            generated_path = generate_resume_docx(
                resume_data_dict=MINIMAL_RESUME_DATA_DICT,
                output_path=str(output_docx_file)
            )
            assert generated_path == str(output_docx_file)
            assert output_docx_file.exists()
            assert output_docx_file.stat().st_size > 0  # Check if file is not empty
        except Exception as e:
            pytest.fail(f"generate_resume_docx raised an exception: {e}")

    def test_basic_content_verification_docx(self, tmp_path):
        """Test basic content of the generated DOCX file."""
        output_docx_file = tmp_path / "test_resume_content.docx"
        
        generate_resume_docx(
            resume_data_dict=MINIMAL_RESUME_DATA_DICT,
            output_path=str(output_docx_file)
        )
        
        assert output_docx_file.exists()
        
        try:
            document = DocxDocument(str(output_docx_file))
            assert len(document.paragraphs) > 5 # Check for a reasonable number of paragraphs

            # Check for name (usually one of the first paragraphs)
            name_found = False
            for para in document.paragraphs[:5]: # Check first 5 paragraphs
                if MINIMAL_RESUME_DATA_DICT["personalInfo"]["name"] in para.text:
                    name_found = True
                    break
            assert name_found, f"Name '{MINIMAL_RESUME_DATA_DICT['personalInfo']['name']}' not found in DOCX."

            # Check for a specific skill (example)
            skills_text_found = False
            skill_to_check = MINIMAL_RESUME_DATA_DICT["technologies_and_skills"][0]
            for para in document.paragraphs:
                if skill_to_check in para.text:
                    skills_text_found = True
                    break
            assert skills_text_found, f"Skill '{skill_to_check}' not found in DOCX."
            
            # Check for an experience company name (usually a heading)
            experience_company_found = False
            company_to_check = MINIMAL_RESUME_DATA_DICT["experience"][0]["company"]
            for para in document.paragraphs: # Headings are also paragraphs
                if company_to_check in para.text:
                    experience_company_found = True
                    break
            assert experience_company_found, f"Company '{company_to_check}' not found in DOCX."

        except Exception as e:
            pytest.fail(f"Error opening or verifying DOCX content: {e}")

    def test_generate_docx_ioerror_on_save_failure(self, tmp_path, monkeypatch):
        """Test that generate_resume_docx raises IOError if document.save() fails."""
        output_docx_file = tmp_path / "locked_resume.docx"

        # Mock document.save() to raise an exception
        def mock_save_raises_exception(*args, **kwargs):
            raise Exception("Simulated save error")

        monkeypatch.setattr("docx.document.Document.save", mock_save_raises_exception)
        
        with pytest.raises(IOError) as excinfo:
            generate_resume_docx(
                resume_data_dict=MINIMAL_RESUME_DATA_DICT,
                output_path=str(output_docx_file)
            )
        assert "Error saving DOCX document" in str(excinfo.value)
        assert "Simulated save error" in str(excinfo.value)
```
