from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_poc_resume_docx(filename="poc_resume.docx"):
    """
    Creates a proof-of-concept DOCX resume using python-docx.
    """
    document = Document()

    # --- Define some basic styles (optional, but good for consistency) ---
    # Attempt to get or create a "Heading1" style for the name
    try:
        name_style = document.styles['Heading 1']
    except KeyError: # Style not found, create a new one based on Normal or use built-in
        name_style = document.styles.add_style('POCNameStyle', WD_STYLE_TYPE.PARAGRAPH)
        name_style.base_style = document.styles['Normal'] # Or another base
        name_style.font.name = 'Calibri'
        name_style.font.size = Pt(24)
        name_style.font.bold = True
        name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_style.paragraph_format.space_after = Pt(12)

    # Attempt to get or create a "Heading2" style for company names
    try:
        company_style = document.styles['Heading 2']
    except KeyError:
        company_style = document.styles.add_style('POCCompanyStyle', WD_STYLE_TYPE.PARAGRAPH)
        company_style.base_style = document.styles['Normal']
        company_style.font.name = 'Calibri'
        company_style.font.size = Pt(14)
        company_style.font.bold = True
        company_style.paragraph_format.space_before = Pt(12)
        company_style.paragraph_format.space_after = Pt(6)

    # --- Sample Data ---
    resume_data = {
        "name": "John Doe",
        "email": "john.doe@example.com",
        "phone": "555-123-4567",
        "experience": [
            {
                "company": "Tech Solutions Inc.",
                "title": "Senior Software Engineer",
                "dates": "2020-Present",
                "details": [
                    "Led development of key features for a flagship product.",
                    "Mentored junior engineers and improved team productivity.",
                    "Collaborated with product managers to define project roadmaps."
                ]
            },
            {
                "company": "Web Innovations LLC",
                "title": "Software Engineer",
                "dates": "2018-2020",
                "details": [
                    "Developed and maintained full-stack web applications.",
                    "Participated in Agile development processes."
                ]
            }
        ],
        "education": {
            "degree": "B.S. in Computer Science",
            "institution": "State University",
            "year": "2018"
        }
    }

    # --- Document Creation ---

    # Name (as a main heading)
    name_paragraph = document.add_paragraph(resume_data["name"], style=name_style)
    # If not using custom/modified built-in style, apply formatting directly:
    # name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # for run in name_paragraph.runs:
    #     run.font.name = 'Calibri' # Example
    #     run.font.size = Pt(24)
    #     run.font.bold = True

    # Contact Info (Email and Phone as a single centered paragraph)
    contact_info = f"{resume_data['email']} | {resume_data['phone']}"
    contact_paragraph = document.add_paragraph(contact_info)
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph.paragraph_format.space_after = Pt(18)


    # Experience Section
    document.add_heading("Experience", level=1) # Built-in Heading 1 style

    for exp in resume_data["experience"]:
        document.add_paragraph(exp["company"], style=company_style) # Using custom/modified style
        
        title_paragraph = document.add_paragraph()
        title_paragraph.add_run(exp["title"]).italic = True
        title_paragraph.add_run(f" ({exp['dates']})")
        title_paragraph.paragraph_format.space_after = Pt(3)

        for detail in exp["details"]:
            document.add_paragraph(detail, style='ListBullet') # Using built-in ListBullet style
    
    document.add_paragraph() # Add some space before next section

    # Education Section
    document.add_heading("Education", level=1)
    edu = resume_data["education"]
    edu_paragraph = document.add_paragraph()
    edu_paragraph.add_run(edu["degree"]).bold = True
    edu_paragraph.add_run(f", {edu['institution']}, {edu['year']}")


    # Save the document
    try:
        document.save(filename)
        print(f"Document '{filename}' created successfully.")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == '__main__':
    create_poc_resume_docx()
